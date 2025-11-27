import pdfplumber
from docx import Document
from fastapi import HTTPException
import requests
import re
from io import BytesIO

class PDFParserService:

    @staticmethod
    def extract_text_from_pdf_source(source: BytesIO) -> str:
        """Extract text from a PDF file-like object"""
        text = ""
        try:
            with pdfplumber.open(source) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text.strip()
        except Exception as e:
            print(f"PDF Extraction Error: {e}")
            return ""


    @staticmethod
    def extract_text_from_pdf(file_path_or_url: str) -> str:
        """Handles both PDFs from local paths and URLs"""
        try:
            # Case 1: URL
            if file_path_or_url.startswith("http://") or file_path_or_url.startswith("https://"):
                response = requests.get(file_path_or_url)
                if response.status_code != 200:
                    raise HTTPException(500, "Failed to download PDF from URL")
                pdf_bytes = BytesIO(response.content)
                return PDFParserService.extract_text_from_pdf_source(pdf_bytes)

            # Case 2: Local file path
            else:
                with open(file_path_or_url, "rb") as f:
                    return PDFParserService.extract_text_from_pdf_source(BytesIO(f.read()))

        except Exception as e:
            raise HTTPException(500, f"Failed to parse PDF: {str(e)}")



    @staticmethod
    def extract_text_from_docx(file_path_or_url: str) -> str:
        """Handles both DOCX from local paths and URLs"""
        try:
            # Case 1: URL
            if file_path_or_url.startswith("http://") or file_path_or_url.startswith("https://"):
                response = requests.get(file_path_or_url)
                if response.status_code != 200:
                    raise HTTPException(500, "Failed to download DOCX from URL")
                doc = Document(BytesIO(response.content))

            # Case 2: Local path
            else:
                doc = Document(file_path_or_url)

            text = ""

            for para in doc.paragraphs:
                text += para.text + "\n"

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"

            if not text.strip():
                raise HTTPException(500, "DOCX file appears empty or unreadable")

            return text.strip()

        except Exception as e:
            raise HTTPException(500, f"Failed to extract text from DOCX: {str(e)}")

    @staticmethod
    def parse_resume(file_path_or_url: str, file_type: str) -> str:
        file_type = file_type.lower()

        if file_type == "pdf":
            return PDFParserService.extract_text_from_pdf(file_path_or_url)

        elif file_type == "docx":
            return PDFParserService.extract_text_from_docx(file_path_or_url)

        else:
            raise HTTPException(400, "Unsupported file type for parsing")


    @staticmethod
    def clean_text(text: str) -> str:
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n+', '\n', text)
        return text.strip()
